#!/usr/bin/env python3
"""Convert CCHP markdown documentation to Word (.docx) files."""

import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
FILES = [
    'PRD_CCHP_Provider_Form_App.md',
    'QRG_Field_Representative.md',
    'QRG_Admin.md',
]

BLUE      = RGBColor(0x1e, 0x4d, 0xb7)
DARK_BLUE = RGBColor(0x0f, 0x32, 0x82)
BLUE_HEX  = '1e4db7'


def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  color_hex)
    tcPr.append(shd)


def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), BLUE_HEX)
    pBdr.append(bot)
    pPr.append(pBdr)


def add_inline(para, text, base_size=None):
    """Parse **bold**, *italic*, `code` and add styled runs to para."""
    # Temporarily replace escaped \* so it doesn't confuse the regex
    text = text.replace(r'\*', '\x00')
    pattern = re.compile(r'\*\*(.+?)\*\*|\*([^*\n]+?)\*|`([^`]+?)`')
    last = 0
    for m in pattern.finditer(text):
        before = text[last:m.start()].replace('\x00', '*')
        if before:
            r = para.add_run(before)
            if base_size:
                r.font.size = Pt(base_size)
        bold_txt, ital_txt, code_txt = m.group(1), m.group(2), m.group(3)
        if bold_txt is not None:
            r = para.add_run(bold_txt.replace('\x00', '*'))
            r.bold = True
            if base_size:
                r.font.size = Pt(base_size)
        elif ital_txt is not None:
            r = para.add_run(ital_txt.replace('\x00', '*'))
            r.italic = True
            if base_size:
                r.font.size = Pt(base_size)
        elif code_txt is not None:
            r = para.add_run(code_txt.replace('\x00', '*'))
            r.font.name = 'Courier New'
            r.font.size = Pt(9.5)
        last = m.end()
    tail = text[last:].replace('\x00', '*')
    if tail:
        r = para.add_run(tail)
        if base_size:
            r.font.size = Pt(base_size)


def strip_inline(text):
    """Remove markdown emphasis markers for plain-text contexts."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*',     r'\1', text)
    text = re.sub(r'`(.+?)`',       r'\1', text)
    return text.replace(r'\*', '*')


def is_separator(cells):
    return all(re.match(r'^[-: ]+$', c) for c in cells if c)


def parse_table(lines, start):
    rows, i = [], start
    while i < len(lines) and lines[i].strip().startswith('|'):
        parts = lines[i].strip().split('|')
        cells = [c.strip() for c in parts[1:-1]] if parts[0] == '' else [c.strip() for c in parts]
        rows.append(cells)
        i += 1
    return rows, i


def convert(md_path, docx_path):
    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.2)
        sec.right_margin  = Inches(1.2)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code = False
    code_buf = []

    while i < len(lines):
        raw     = lines[i].rstrip('\n')
        stripped = raw.strip()

        # ── Code fence ─────────────────────────────────────────────────────
        if stripped.startswith('```'):
            if not in_code:
                in_code, code_buf = True, []
            else:
                in_code = False
                if code_buf:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent  = Inches(0.4)
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after  = Pt(4)
                    r = p.add_run('\n'.join(code_buf))
                    r.font.name  = 'Courier New'
                    r.font.size  = Pt(9)
                    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            i += 1
            continue

        if in_code:
            code_buf.append(raw)
            i += 1
            continue

        # ── Blank line ──────────────────────────────────────────────────────
        if stripped == '':
            i += 1
            continue

        # ── Horizontal rule ─────────────────────────────────────────────────
        if stripped in ('---', '***', '___') and len(stripped) >= 3 and len(set(stripped)) == 1:
            add_hr(doc)
            i += 1
            continue

        # ── Heading 1 ───────────────────────────────────────────────────────
        if re.match(r'^# [^#]', raw):
            h = doc.add_heading(raw[2:].strip(), level=1)
            if h.runs:
                h.runs[0].font.color.rgb = DARK_BLUE
            i += 1
            continue

        # ── Heading 2 ───────────────────────────────────────────────────────
        if re.match(r'^## [^#]', raw):
            h = doc.add_heading(raw[3:].strip(), level=2)
            if h.runs:
                h.runs[0].font.color.rgb = BLUE
            i += 1
            continue

        # ── Heading 3 ───────────────────────────────────────────────────────
        if re.match(r'^### [^#]', raw):
            h = doc.add_heading(raw[4:].strip(), level=3)
            if h.runs:
                h.runs[0].font.color.rgb = DARK_BLUE
            i += 1
            continue

        # ── Table ────────────────────────────────────────────────────────────
        if stripped.startswith('|'):
            all_rows, i = parse_table(lines, i)
            data = [r for r in all_rows if not is_separator(r)]
            if not data:
                continue
            ncols = max(len(r) for r in data)
            tbl   = doc.add_table(rows=len(data), cols=ncols)
            tbl.style = 'Table Grid'
            for ri, row in enumerate(data):
                for ci in range(ncols):
                    txt  = row[ci] if ci < len(row) else ''
                    cell = tbl.rows[ri].cells[ci]
                    p    = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after  = Pt(2)
                    if ri == 0:
                        set_cell_bg(cell, BLUE_HEX)
                        r = p.add_run(strip_inline(txt))
                        r.bold = True
                        r.font.size      = Pt(10)
                        r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                    else:
                        add_inline(p, txt, base_size=10)
            doc.add_paragraph()
            continue

        # ── Blockquote ───────────────────────────────────────────────────────
        if stripped.startswith('>'):
            content = re.sub(r'^>\s*', '', stripped)
            if re.match(r'^[-*]\s', content):
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.5)
                add_inline(p, content[2:])
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent  = Inches(0.4)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
                add_inline(p, content)
                for r in p.runs:
                    r.font.color.rgb = BLUE
                    r.italic = True
            i += 1
            continue

        # ── Unordered list ───────────────────────────────────────────────────
        ul = re.match(r'^(\s*)[-*]\s+(.+)', raw)
        if ul:
            indent, content = len(ul.group(1)), ul.group(2)
            p = doc.add_paragraph(style='List Bullet')
            if indent >= 2:
                p.paragraph_format.left_indent = Inches(0.5)
            add_inline(p, content)
            i += 1
            continue

        # ── Ordered list ─────────────────────────────────────────────────────
        ol = re.match(r'^\d+\.\s+(.+)', raw)
        if ol:
            p = doc.add_paragraph(style='List Number')
            add_inline(p, ol.group(1))
            i += 1
            continue

        # ── Plain paragraph ───────────────────────────────────────────────────
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f'  saved → {os.path.basename(docx_path)}')


if __name__ == '__main__':
    print('Converting CCHP markdown to Word...')
    for fname in FILES:
        src = os.path.join(DOCS_DIR, fname)
        dst = os.path.join(DOCS_DIR, fname.replace('.md', '.docx'))
        print(f'  {fname}')
        convert(src, dst)
    print('Done!')
